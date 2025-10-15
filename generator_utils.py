# ==================================================
# File: generator_utils.py
# Générateur de rapports - Simple et IA
# ==================================================

import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai
import os
from multilingual_config import MultilingualConfig
from typing import Dict, List, Any
import io

class ReportGenerator:
    """Générateur de rapports démographiques"""
    
    def __init__(self, ml_config: MultilingualConfig):
        self.ml_config = ml_config
        self.gemini_api_key = os.getenv("GEMINI_API_KEY", "")
        
        if self.gemini_api_key:
            genai.configure(api_key=self.gemini_api_key)
            self.model = genai.GenerativeModel('gemini-2.0-flash-exp')
    
    def create_simple_report(self, module_name: str, data: Dict[str, Any], figures: List = None) -> bytes:
        """Génère rapport Word simple"""
        
        doc = Document()
        self._add_header(doc, module_name)
        
        date_text = datetime.now().strftime("%d/%m/%Y" if self.ml_config.get_language() == "fr" else "%m/%d/%Y")
        doc.add_paragraph(f"{'Date' if self.ml_config.get_language() == 'en' else 'Date'}: {date_text}")
        doc.add_paragraph()
        
        self._add_executive_summary(doc, module_name, data)
        self._add_main_data(doc, data)
        
        if figures:
            self._add_figures(doc, figures)
        
        self._add_brief_explanation(doc, module_name, data)
        self._add_footer(doc)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def create_ai_report(self, module_name: str, data: Dict[str, Any], figures: List = None) -> bytes:
        """Génère rapport Word avec analyse IA"""
        
        doc = Document()
        self._add_header(doc, module_name)
        
        date_text = datetime.now().strftime("%d/%m/%Y" if self.ml_config.get_language() == "fr" else "%m/%d/%Y")
        doc.add_paragraph(f"{'Date' if self.ml_config.get_language() == 'en' else 'Date'}: {date_text}")
        doc.add_paragraph()
        
        self._add_executive_summary(doc, module_name, data)
        self._add_main_data(doc, data)
        
        if figures:
            self._add_figures(doc, figures)
        
        self._add_ai_analysis(doc, module_name, data)
        self._add_footer(doc)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _add_header(self, doc: Document, module_name: str):
        title = doc.add_heading('Africa Demographics Platform', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle_text = f"{'Rapport' if self.ml_config.get_language() == 'fr' else 'Report'}: {module_name}"
        subtitle = doc.add_heading(subtitle_text, level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    def _add_executive_summary(self, doc: Document, module_name: str, data: Dict[str, Any]):
        heading_text = "Résumé Exécutif" if self.ml_config.get_language() == "fr" else "Executive Summary"
        doc.add_heading(heading_text, level=2)
        
        summary_lines = []
        
        if 'total_population_millions' in data:
            pop = data['total_population_millions']
            summary_lines.append(f"Population totale: {pop:.1f} millions")
        
        if 'weighted_tfr' in data and not pd.isna(data['weighted_tfr']):
            tfr = data['weighted_tfr']
            summary_lines.append(f"Taux de fécondité moyen: {tfr:.2f}")
        
        if 'weighted_median_age' in data and not pd.isna(data['weighted_median_age']):
            age = data['weighted_median_age']
            summary_lines.append(f"Âge médian: {age:.1f} ans")
        
        if 'countries_analyzed' in data:
            countries = data['countries_analyzed']
            summary_lines.append(f"Pays analysés: {countries}")
        
        for line in summary_lines:
            doc.add_paragraph(line, style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_main_data(self, doc: Document, data: Dict[str, Any]):
        heading_text = "Données Principales" if self.ml_config.get_language() == "fr" else "Main Data"
        doc.add_heading(heading_text, level=2)
        
        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, (int, float)) and not isinstance(value, bool):
                    if not pd.isna(value):
                        p = doc.add_paragraph()
                        p.add_run(f"{key}: ").bold = True
                        p.add_run(f"{value:.2f}" if isinstance(value, float) else str(value))
                elif isinstance(value, str):
                    p = doc.add_paragraph()
                    p.add_run(f"{key}: ").bold = True
                    p.add_run(str(value))
        
        doc.add_paragraph()
    
    def _add_figures(self, doc: Document, figures: List):
        heading_text = "Graphiques" if self.ml_config.get_language() == "fr" else "Figures"
        doc.add_heading(heading_text, level=2)
        
        for i, fig in enumerate(figures):
            if fig is None:
                continue
                
            try:
                img_bytes = fig.to_image(format="png", width=800, height=600)
                image_stream = io.BytesIO(img_bytes)
                doc.add_picture(image_stream, width=Inches(6))
                
                caption_text = f"Figure {i+1}"
                caption = doc.add_paragraph(caption_text)
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            except Exception as e:
                doc.add_paragraph(f"[Erreur figure {i+1}: {str(e)}]")
        
        doc.add_paragraph()
    
    def _add_brief_explanation(self, doc: Document, module_name: str, data: Dict[str, Any]):
        heading_text = "Explication" if self.ml_config.get_language() == "fr" else "Explanation"
        doc.add_heading(heading_text, level=2)
        
        explanations = {
            "Vue Continentale": "Ce rapport présente une vue d'ensemble des indicateurs démographiques pour l'ensemble du continent africain, calculés par moyennes pondérées par population.",
            "Continental Overview": "This report presents an overview of demographic indicators for the entire African continent, calculated using population-weighted averages.",
            "Profils Pays": "Ce rapport analyse en détail les indicateurs démographiques pour le pays sélectionné, incluant la structure par âge et les tendances temporelles.",
            "Country Profiles": "This report analyzes in detail the demographic indicators for the selected country, including age structure and temporal trends.",
            "Analyse des Tendances": "Ce rapport compare l'évolution temporelle des indicateurs démographiques entre plusieurs pays.",
            "Trend Analysis": "This report compares the temporal evolution of demographic indicators across multiple countries.",
            "Analyse par Groupement": "Ce rapport utilise des algorithmes de machine learning pour classifier les pays africains selon leurs profils démographiques similaires.",
            "Clustering Analysis": "This report uses machine learning algorithms to classify African countries according to their similar demographic profiles.",
            "Explorateur de Données": "Ce rapport présente les données filtrées selon vos critères de sélection.",
            "Data Explorer": "This report presents the filtered data according to your selection criteria."
        }
        
        explanation = explanations.get(module_name, "Ce rapport présente une analyse démographique détaillée.")
        doc.add_paragraph(explanation)
        doc.add_paragraph()
    
    def _add_ai_analysis(self, doc: Document, module_name: str, data: Dict[str, Any]):
        heading_text = "Analyse IA et Recommandations" if self.ml_config.get_language() == "fr" else "AI Analysis and Recommendations"
        doc.add_heading(heading_text, level=2)
        
        if not self.gemini_api_key:
            doc.add_paragraph("⚠️ Clé API Gemini non configurée. Définissez GEMINI_API_KEY.")
            return
        
        prompt = self._generate_ai_prompt(module_name, data)
        
        try:
            response = self.model.generate_content(prompt)
            analysis = response.text
            doc.add_paragraph(analysis)
            
        except Exception as e:
            error_text = f"Erreur analyse IA: {str(e)}" if self.ml_config.get_language() == "fr" else f"Error during AI analysis: {str(e)}"
            doc.add_paragraph(error_text)
        
        doc.add_paragraph()
    
    def _generate_ai_prompt(self, module_name: str, data: Dict[str, Any]) -> str:
        lang = self.ml_config.get_language()
        
        if lang == "fr":
            prompt = f"""Tu es un expert en démographie africaine. Analyse les données suivantes du module "{module_name}" et fournis:

1. **Interprétation des Résultats**: Explique ces données de manière simple et accessible.

2. **Lecture Simplifiée**: Résume les points clés en 3-5 bullets pour des non-experts.

3. **Recommandations Politiques**: Propose 3-5 recommandations concrètes pour:
   - Politiques de santé publique
   - Éducation et formation
   - Emploi et économie
   - Planning familial

Données:
{self._format_data_for_prompt(data)}

Réponds en français, de manière claire et actionable."""
        else:
            prompt = f"""You are an expert in African demography. Analyze the following data from "{module_name}" and provide:

1. **Results Interpretation**: Explain this data in a simple and accessible way.

2. **Simplified Reading**: Summarize key points in 3-5 bullets for non-experts.

3. **Policy Recommendations**: Propose 3-5 concrete recommendations for:
   - Public health policies
   - Education and training
   - Employment and economy
   - Family planning

Data:
{self._format_data_for_prompt(data)}

Respond in English, clearly and actionably."""
        
        return prompt
    
    def _format_data_for_prompt(self, data: Dict[str, Any]) -> str:
        formatted = []
        
        for key, value in data.items():
            if isinstance(value, dict):
                formatted.append(f"{key}:")
                for sub_key, sub_value in value.items():
                    if isinstance(sub_value, (int, float)) and not pd.isna(sub_value):
                        formatted.append(f"  - {sub_key}: {sub_value:.2f}")
                    elif not pd.isna(sub_value):
                        formatted.append(f"  - {sub_key}: {sub_value}")
            elif isinstance(value, (int, float)) and not pd.isna(value):
                formatted.append(f"{key}: {value:.2f}")
            elif isinstance(value, str):
                formatted.append(f"{key}: {value}")
        
        return "\n".join(formatted)
    
    def _add_footer(self, doc: Document):
        doc.add_paragraph()
        doc.add_paragraph("─" * 80)
        
        footer_text = f"""
Africa Demographics Platform (ADP) v2.5
Conception et développement: Zakaria Benhoumad
Assisté par: Anthropic Claude
Source des données: World Bank Open Data API
Généré le: {datetime.now().strftime("%d/%m/%Y %H:%M")}
        """
        
        footer = doc.add_paragraph(footer_text.strip())
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in footer.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)


def add_report_buttons(module_name: str, data: Dict[str, Any], figures: List = None, ml_config: MultilingualConfig = None):
    """Ajoute boutons de génération de rapports"""
    
    if ml_config is None or not data:
        return
    
    st.markdown("---")
    
    report_title = "📄 Génération de Rapport" if ml_config.get_language() == "fr" else "📄 Report Generation"
    st.markdown(f"### {report_title}")
    
    col1, col2 = st.columns(2)
    
    generator = ReportGenerator(ml_config)
    
    with col1:
        simple_text = "📝 Générer Rapport Simple" if ml_config.get_language() == "fr" else "📝 Generate Simple Report"
        if st.button(simple_text, key=f"simple_report_{module_name}"):
            with st.spinner("Génération..." if ml_config.get_language() == "fr" else "Generating..."):
                try:
                    report_bytes = generator.create_simple_report(module_name, data, figures)
                    
                    filename = f"rapport_simple_{module_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    
                    st.download_button(
                        label="⬇️ Télécharger Rapport" if ml_config.get_language() == "fr" else "⬇️ Download Report",
                        data=report_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_simple_{module_name}"
                    )
                    
                    st.success("✅ Rapport généré!" if ml_config.get_language() == "fr" else "✅ Report generated!")
                    
                except Exception as e:
                    st.error(f"Erreur: {str(e)}")
    
    with col2:
        ai_text = "🤖 Générer Rapport via IA (Gemini)" if ml_config.get_language() == "fr" else "🤖 Generate AI Report (Gemini)"
        if st.button(ai_text, key=f"ai_report_{module_name}"):
            if not generator.gemini_api_key:
                st.warning("⚠️ Configurez GEMINI_API_KEY pour utiliser cette fonctionnalité.")
                return
                
            with st.spinner("Analyse IA..." if ml_config.get_language() == "fr" else "AI analysis..."):
                try:
                    report_bytes = generator.create_ai_report(module_name, data, figures)
                    
                    filename = f"rapport_ia_{module_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    
                    st.download_button(
                        label="⬇️ Télécharger Rapport IA" if ml_config.get_language() == "fr" else "⬇️ Download AI Report",
                        data=report_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_ai_{module_name}"
                    )
                    
                    st.success("✅ Rapport IA généré!" if ml_config.get_language() == "fr" else "✅ AI report generated!")
                    
                except Exception as e:
                    st.error(f"Erreur: {str(e)}")
    
    if ml_config.get_language() == "fr":
        st.info("💡 Configurez GEMINI_API_KEY pour activer le rapport IA.")
    else:
        st.info("💡 Configure GEMINI_API_KEY to enable AI report.")